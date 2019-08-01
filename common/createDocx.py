#-*- coding:utf-8 -*-

import sys 
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH # 정렬하는 방법
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Inches, Pt

# Global variable
document = None # For making a docx file
user_information = None # user model
user_time = None    # workTime model
tableSize = None    # size of workTime model


def __get_user_information__(user_info, userTime, loopSize):
    global user_information
    global user_time
    global tableSize

    user_information = user_info
    user_time = userTime
    tableSize = loopSize
    tableSize = int(tableSize)     



def __init_the_Document__():
    global document 
    document = Document()
    style = document.styles['Normal']
    font = style.font
    # font.name = 'Arial'
    font.size = Pt(10)



def __set_table_header(user_information, date_information):
    heading = None
    heading_paragraph = None
    table_head_1 = None
    table_head_2 = None
    table_head_1_row = None
    table_head_1_str = None
    user_info = None
    date_info = None
    global document


    heading = document.add_heading('봉사장학생 근무확인표')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER # 가운데 정렬하는 방법
    # 같은 의미:: paragraph.alignment = 0  # for left, 1 for right, 2 center, 3 justify ....


    date_info = date_information.split(',')
    date_string = "("
    date_string += date_info[0]
    date_string += '학년도 '
    date_string += date_info[1]
    date_string += '월분)'

    heading_paragraph = document.add_paragraph(date_string)
    heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT # 왼쪽 정렬하는 방법


    # make the table architecture
    table_head_1 = document.add_table(rows = 1, cols = 1, style='Table Grid')
    table_head_2 = document.add_table(rows = 2, cols = 4, style='Table Grid')

    # make the table alignment
    table_head_1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_head_2.alignment = WD_TABLE_ALIGNMENT.CENTER


    # push contents at table 
    user_info = user_information.split(',')
    
    # push contents at table_head_1
    table_head_1_row = table_head_1.rows[0]

    table_head_1_str = "홍익대학 "
    table_head_1_str += "  학부 "
    table_head_1_str += user_info[2]
    table_head_1_str += " "
    table_head_1_str += user_info[3]
    table_head_1_str += "학년"
    table_head_1_row.cells[0].text = table_head_1_str

    # push contents at table_head_2

    table_head_2.cell(0, 0).text = '성 명'
    table_head_2.cell(1, 0).text = '근무부서'
    table_head_2.cell(1, 1).text = '취업진로지원센터'
    table_head_2.cell(0, 2).text = '학 번'
    table_head_2.cell(1, 2).text = '봉사분야'
    
    table_head_2.cell(0, 1).text = user_info[1]
    table_head_2.cell(0, 3).text = user_info[0]
    table_head_2.cell(1, 3).text = user_info[5]



def __set_table_body(table_row_size, user_time):
    table_body = None
    table_body_cells = None
    table_body_text = None
    user_workTime = None
    BODY_ROWSIZE = None
    LOOPSIZE = None
    offset = None
    global document

    # temp variable
    table_body = document.add_table(rows = table_row_size + 1, cols = 9, style='Table Grid')
    table_body.alignment = WD_TABLE_ALIGNMENT.CENTER

    # push contents adjust alignment at table_body
    table_body_text = ['근무월일', '요일', '근무시간', '근로 상세내역', '계', '누 계', '학생 확인', '담당자 확인', '부서장 확인']

    for i in range(len(table_body_text)):
        table_body.cell(0, i).text = table_body_text[i]
        table_body.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    table_body_cells = table_body.rows[0].cells
    table_body_cells[2].width = Cm(6)
    table_body_cells[3].width = Cm(6)

    # push contents at table_body (user_workTime)
    user_workTime = user_time.split(',')
    BODY_ROWSIZE = 10
    LOOPSIZE = len(user_workTime) // BODY_ROWSIZE

    offset = 0
    for i in range(LOOPSIZE):
        table_body_contents = [] 
        time_line_string = ""           
        for j in range(BODY_ROWSIZE):
            if (j == 2):
                time_line_string += user_workTime[offset]
                time_line_string += " ~ "
            elif (j == 3):
                time_line_string += user_workTime[offset]
                table_body_contents.append(time_line_string)
            else:
                table_body_contents.append(user_workTime[offset])
            
            offset += 1

        for idx in range(len(table_body_contents)):
            table_body.cell(i+1, idx).text = table_body_contents[idx]
            table_body.cell(i+1, idx).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    


def __set_table_tail():
    table_tail_1 = None
    table_tail_2 = None
    table_tail_1_text = None
    table_tail_2_cells = None
    global document

    table_tail_1 = document.add_table(rows = 1, cols = 3, style='Table Grid')
    table_tail_2 = document.add_table(rows = 1, cols = 1, style='Table Grid')
    
    table_tail_1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_tail_2.alignment = WD_TABLE_ALIGNMENT.CENTER

    # push contents at table_tail_1
    table_tail_1_text = ['기준근무시간 : □85시간  □44시간  □42시간', '총 근무시간    시간   분 (인)', '근무평정    초 과 · 부 족 (   시간    분)']

    for i in range(len(table_tail_1_text)):
        table_tail_1.cell(0, i).text = table_tail_1_text[i]
        table_tail_1.cell(0, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # push contents at table_tail_2
    table_tail_2_cells = table_tail_2.rows[0].cells
    p = table_tail_2_cells[0].add_paragraph('1. 봉사장학생은 근무부서 부서장의 지시에 따라 성실히 근무에 임하여야 하며, 근무성적이 불량하거나 지시에 불응하여 계속 근무가 곤란하다고 판단 될 때에는 봉사장학생 자격을 취소할 수 있습니다.')
    p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p = table_tail_2_cells[0].add_paragraph('2. 월정액(월355,000원) 봉사장학생은 근무시간이 월 42시간 기준이며, 42시간을 근무하지 않은 경우에는 봉사 장학금이 지급되지 않고, 시간급일 경우 월85시간(시간3급,4급은 월44시간)을 초과하여 근무할 수 없습니다.')
    p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p = table_tail_2_cells[0].add_paragraph('3. 봉사장학금은 학생의 클래스넷 개인정보에 입력된 계좌번호로 매월 10일 입금합니다.')
    p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p = table_tail_2_cells[0].add_paragraph('4. 각 부서장은 봉사장학생의 대리근무 여부 및 재학생 이외의 부적격자(휴학, 제적, 졸업생 등)가 근로를 하는 일이 없도록 관리를 철저히 하여 주시기 바랍니다.')
    p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p = table_tail_2_cells[0].add_paragraph('5. 월별 봉사장학생 근무확인표 원본은 해당부서에서 보관하시기 바랍니다.')
    p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p = table_tail_2_cells[0].add_paragraph('6. 봉사장학생 근무확인표는 근무부서에서 수정하여 사용할 수 있습니다.(근무시간, 학생-담당자-부서장 확인 필수)')
    p.alignment=WD_ALIGN_PARAGRAPH.LEFT



def __save_the_Document__(user_information, path):
    user_info = None
    docx_file_name = path + ""
    docx_file_name += "/uploads/"
    global document

    user_info = user_information.split(',')
    docx_file_name += user_info[0]  #학번.docx
    docx_file_name += ".docx"

    document.save(docx_file_name)



'''
## Main Logic ##
'''
__get_user_information__(sys.argv[1], sys.argv[2], sys.argv[3])
__init_the_Document__()
__set_table_header(user_information, sys.argv[4])
__set_table_body(tableSize, user_time)
__set_table_tail()
__save_the_Document__(user_information, sys.argv[5])

# for response:: this is a dummy response
print('success')
